"""Render, validate, and publish complete syllabus artifacts."""

# Standard Library
import os
import re
import sys
import html
import shutil
import pathlib
import zipfile
import subprocess
import importlib.util

# PIP3 modules
import docx
import pypdf
import markdown
import docx.enum.text
import docx.oxml
import docx.oxml.ns
import docx.shared

# local repo modules
import build_lib.syllabus_model
import build_lib.syllabus_content


# Generated public artifacts are documents, never executable server-side content (ASVS 5.3.1).
MANAGED_DOWNLOAD_SUFFIXES = {".docx", ".pdf"}


#============================================
def mark_table_header(row: object) -> None:
	"""Mark one DOCX row as a repeating semantic table header."""
	properties = row._tr.get_or_add_trPr()
	if properties.find(docx.oxml.ns.qn("w:tblHeader")) is None:
		header = docx.oxml.OxmlElement("w:tblHeader")
		header.set(docx.oxml.ns.qn("w:val"), "true")
		properties.append(header)
	return None


#============================================
def format_table_header(row: object) -> None:
	"""Make table headers distinct through both bold text and light shading."""
	mark_table_header(row)
	for cell in row.cells:
		cell_properties = cell._tc.get_or_add_tcPr()
		shading = cell_properties.find(docx.oxml.ns.qn("w:shd"))
		if shading is None:
			shading = docx.oxml.OxmlElement("w:shd")
			cell_properties.append(shading)
		shading.set(docx.oxml.ns.qn("w:fill"), "E6E6E6")
		for paragraph in cell.paragraphs:
			for run in paragraph.runs:
				run.bold = True
	return None


#============================================
def prevent_row_split(row: object) -> None:
	"""Keep a DOCX table row from splitting across pages."""
	properties = row._tr.get_or_add_trPr()
	if properties.find(docx.oxml.ns.qn("w:cantSplit")) is None:
		cant_split = docx.oxml.OxmlElement("w:cantSplit")
		properties.append(cant_split)
	return None


#============================================
def set_document_language(document: object, language_code: str) -> None:
	"""Set the proofing language on paragraph and character styles."""
	for style in document.styles:
		run_properties = style.element.get_or_add_rPr()
		language = run_properties.find(docx.oxml.ns.qn("w:lang"))
		if language is None:
			language = docx.oxml.OxmlElement("w:lang")
			run_properties.append(language)
		language.set(docx.oxml.ns.qn("w:val"), language_code)
	return None


#============================================
def append_word_field(paragraph: object, instruction_text: str, fallback_text: str) -> None:
	"""Append one standard Word field with a readable fallback value."""
	run = paragraph.add_run()
	field_begin = docx.oxml.OxmlElement("w:fldChar")
	field_begin.set(docx.oxml.ns.qn("w:fldCharType"), "begin")
	instruction = docx.oxml.OxmlElement("w:instrText")
	instruction.set(docx.oxml.ns.qn("xml:space"), "preserve")
	instruction.text = f" {instruction_text} "
	field_separator = docx.oxml.OxmlElement("w:fldChar")
	field_separator.set(docx.oxml.ns.qn("w:fldCharType"), "separate")
	fallback = docx.oxml.OxmlElement("w:t")
	fallback.text = fallback_text
	field_end = docx.oxml.OxmlElement("w:fldChar")
	field_end.set(docx.oxml.ns.qn("w:fldCharType"), "end")
	for element in (field_begin, instruction, field_separator, fallback, field_end):
		run._r.append(element)
	return None


#============================================
def configure_docx_footer(
	document: object,
	manifest: build_lib.syllabus_model.SyllabusManifest,
) -> None:
	"""Add course, current-section, and Page X of Y fields to every footer."""
	for section in document.sections:
		footer = section.footer
		footer.paragraphs[0].clear()
		footer.paragraphs[0].paragraph_format.space_after = docx.shared.Pt(0)
		usable_width = section.page_width - section.left_margin - section.right_margin
		table = footer.add_table(rows=1, cols=3, width=usable_width)
		table.autofit = False
		cell_width = usable_width // 3
		for cell in table.rows[0].cells:
			cell.width = cell_width
		borders = docx.oxml.OxmlElement("w:tblBorders")
		for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
			border = docx.oxml.OxmlElement(f"w:{edge}")
			border.set(docx.oxml.ns.qn("w:val"), "nil")
			borders.append(border)
		table._tbl.tblPr.append(borders)
		left, center, right = [cell.paragraphs[0] for cell in table.rows[0].cells]
		left.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
		left.add_run(f"{manifest.course_code} - {manifest.term}")
		center.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
		append_word_field(center, 'STYLEREF "Heading 2"', "Current section")
		right.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT
		right.add_run("Page ")
		append_word_field(right, "PAGE", "1")
		right.add_run(" of ")
		append_word_field(right, "NUMPAGES", "1")
		for paragraph in (left, center, right):
			paragraph.paragraph_format.space_after = docx.shared.Pt(0)
			for run in paragraph.runs:
				run.font.name = "Arial"
				run.font.size = docx.shared.Pt(8)
	settings = document.settings.element
	update_fields = settings.find(docx.oxml.ns.qn("w:updateFields"))
	if update_fields is None:
		update_fields = docx.oxml.OxmlElement("w:updateFields")
		settings.append(update_fields)
	update_fields.set(docx.oxml.ns.qn("w:val"), "true")
	return None


#============================================
def configure_docx_contents(document: object) -> None:
	"""Add dotted leaders and page-reference fields to the linked contents list."""
	contents_started = False
	# LibreOffice cannot resolve these two Pandoc section bookmarks as page references. Their first
	# subsection bookmarks occupy the same starting page and remain interoperable with Word.
	page_reference_aliases = {
		"course-details": "course-overview",
		"student-resources": "academic-progress-and-learning",
	}
	usable_width = (
		document.sections[0].page_width
		- document.sections[0].left_margin
		- document.sections[0].right_margin
	)
	for paragraph in document.paragraphs:
		if paragraph.style.name == "Heading 1" and paragraph.text.strip() == "Contents":
			contents_started = True
			continue
		if not contents_started:
			continue
		if paragraph.style.name.startswith("Heading "):
			break
		hyperlink = paragraph._p.find(docx.oxml.ns.qn("w:hyperlink"))
		if hyperlink is None:
			continue
		anchor = hyperlink.get(docx.oxml.ns.qn("w:anchor"))
		if not anchor:
			continue
		page_anchor = page_reference_aliases.get(anchor, anchor)
		paragraph.paragraph_format.tab_stops.add_tab_stop(
			usable_width,
			docx.enum.text.WD_TAB_ALIGNMENT.RIGHT,
			docx.enum.text.WD_TAB_LEADER.DOTS,
		)
		paragraph.add_run("\t")
		append_word_field(paragraph, f"PAGEREF {page_anchor} \\h", "1")
	return None


#============================================
def postprocess_docx(
	docx_path: pathlib.Path,
	manifest: build_lib.syllabus_model.SyllabusManifest,
) -> None:
	"""Apply metadata, language, and table accessibility properties."""
	document = docx.Document(docx_path)
	document.core_properties.title = f"{manifest.course_code}: {manifest.title} - {manifest.term}"
	document.core_properties.author = manifest.author
	document.core_properties.subject = "Complete course syllabus"
	document.core_properties.language = manifest.language
	set_document_language(document, manifest.language)
	configure_docx_footer(document, manifest)
	configure_docx_contents(document)
	for table in document.tables:
		table.style = "Table"
		table.autofit = True
		if table.rows:
			format_table_header(table.rows[0])
		for row in table.rows:
			prevent_row_split(row)
	document.save(docx_path)
	return None


#============================================
def audit_docx_structure(
	docx_path: pathlib.Path,
	manifest: build_lib.syllabus_model.SyllabusManifest,
) -> None:
	"""Report non-blocking DOCX accessibility findings."""
	document = docx.Document(docx_path)
	accessibility_findings = []
	expected_title = f"{manifest.course_code}: {manifest.title} - {manifest.term}"
	if document.core_properties.title != expected_title:
		accessibility_findings.append("missing expected document title metadata")
	if document.core_properties.language != manifest.language:
		accessibility_findings.append("missing expected document language metadata")
	heading_paragraphs = [
		paragraph
		for paragraph in document.paragraphs
		if paragraph.style.name.startswith("Heading ")
	]
	if not heading_paragraphs:
		accessibility_findings.append("no semantic heading styles found")
	if not document.tables:
		accessibility_findings.append("no semantic tables found")
	for table in document.tables:
		header_properties = table.rows[0]._tr.get_or_add_trPr()
		if header_properties.find(docx.oxml.ns.qn("w:tblHeader")) is None:
			accessibility_findings.append("table is missing a repeating header row")
	for finding in accessibility_findings:
		print(f"Accessibility advisory: {docx_path}: {finding}")
	return None


#============================================
def verify_docx_output(
	docx_path: pathlib.Path,
	manifest: build_lib.syllabus_model.SyllabusManifest,
) -> None:
	"""Require a readable DOCX package containing every manifest section."""
	document = docx.Document(docx_path)
	paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
	build_lib.syllabus_content.verify_required_section_titles(
		paragraph_text,
		manifest,
		docx_path,
	)
	combined_markdown = build_lib.syllabus_content.compose_markdown(manifest)
	build_lib.syllabus_content.validate_markdown_tables(combined_markdown, str(manifest.path))
	expected_table_count = build_lib.syllabus_content.count_markdown_tables(combined_markdown)
	if len(document.tables) != expected_table_count:
		raise RuntimeError(
			f"{docx_path}: expected {expected_table_count} tables, found {len(document.tables)}"
		)
	return None


#============================================
def run_pandoc_docx(
	markdown_path: pathlib.Path,
	docx_path: pathlib.Path,
	reference_path: pathlib.Path,
	manifest: build_lib.syllabus_model.SyllabusManifest,
) -> None:
	"""Generate one complete DOCX through Pandoc."""
	# ASVS 1.2.5: pass trusted paths and metadata as an argument list without a shell.
	command = [
		"pandoc",
		str(markdown_path),
		"--from=markdown+pipe_tables",
		"--to=docx",
		f"--reference-doc={reference_path}",
		f"--metadata=title:{manifest.course_code}: {manifest.title}",
		f"--metadata=subtitle:{manifest.term}",
		f"--metadata=author:{manifest.author}",
		f"--metadata=lang:{manifest.language}",
		f"--output={docx_path}",
	]
	subprocess.run(command, check=True)
	return None


#============================================
def run_markdown_html(
	markdown_text: str,
	html_path: pathlib.Path,
	stylesheet_path: pathlib.Path,
	manifest: build_lib.syllabus_model.SyllabusManifest,
	extensions: tuple[str, ...],
	extension_configs: dict[str, dict[object, object]],
) -> None:
	"""Render semantic standalone HTML with the site's Markdown extension stack."""
	build_lib.syllabus_content.validate_markdown_tables(markdown_text, str(manifest.path))
	converter = markdown.Markdown(
		extensions=list(extensions),
		extension_configs=extension_configs,
		output_format="html5",
	)
	body_html = converter.convert(markdown_text)
	expected_table_count = build_lib.syllabus_content.count_markdown_tables(markdown_text)
	rendered_table_count = body_html.count("<table>")
	if rendered_table_count != expected_table_count:
		raise RuntimeError(
			f"{html_path}: expected {expected_table_count} tables, found {rendered_table_count}"
		)
	document_title = f"{manifest.course_code}: {manifest.title} - {manifest.term}"
	# ASVS 1.1.2 and 1.2.1: escape metadata at the final HTML rendering boundary.
	escaped_title = html.escape(document_title)
	escaped_course_title = html.escape(f"{manifest.course_code}: {manifest.title}")
	escaped_term = html.escape(manifest.term)
	escaped_author = html.escape(manifest.author)
	escaped_language = html.escape(manifest.language, quote=True)
	# ASVS 1.1.2 and 1.2.1: escape the allowlisted CSS token at the HTML boundary.
	escaped_course_color = html.escape(manifest.course_color, quote=True)
	document = "\n".join(
		[
			"<!doctype html>",
			f'<html lang="{escaped_language}">',
			"<head>",
			'<meta charset="utf-8">',
			f"<title>{escaped_title}</title>",
			f'<meta name="author" content="{escaped_author}">',
			f'<link rel="stylesheet" href="{stylesheet_path.as_uri()}">',
			"</head>",
			(
				'<body class="syllabus-document" '
				f'style="--syllabus-page-accent: {escaped_course_color}">'
			),
			'<header id="title-block-header">',
			f"<h1>{escaped_course_title}</h1>",
			f'<p class="subtitle">{escaped_term}</p>',
			f'<p class="author">{escaped_author}</p>',
			"</header>",
			'<main id="syllabus-content">',
			body_html,
			"</main>",
			"</body>",
			"</html>",
			"",
		]
	)
	html_path.write_text(document, encoding="utf-8")
	return None


#============================================
def run_weasyprint_pdf(
	html_path: pathlib.Path,
	pdf_path: pathlib.Path,
) -> None:
	"""Render standalone syllabus HTML directly to a tagged PDF."""
	# ASVS 1.2.5: pass trusted staged artifact paths without shell interpolation.
	command = [
		sys.executable,
		"-m",
		"weasyprint",
		"--pdf-tags",
		str(html_path),
		str(pdf_path),
	]
	subprocess.run(command, check=True)
	if not pdf_path.is_file():
		raise RuntimeError(f"WeasyPrint did not create {pdf_path}")
	return None


#============================================
def verify_pdf_output(
	pdf_path: pathlib.Path,
	manifest: build_lib.syllabus_model.SyllabusManifest,
) -> None:
	"""Require a usable PDF and report non-blocking accessibility findings."""
	# ASVS 1.2.5: invoke fixed local tools with a validated generated artifact path.
	metadata = subprocess.run(
		["pdfinfo", str(pdf_path)],
		check=True,
		capture_output=True,
		text=True,
	).stdout
	if re.search(r"^Page size:\s+612 x 792 pts", metadata, re.MULTILINE) is None:
		raise RuntimeError(f"{pdf_path}: PDF is not US letter size")
	reader = pypdf.PdfReader(pdf_path)
	if not reader.pages:
		raise RuntimeError(f"{pdf_path}: PDF contains no pages")
	text_content = subprocess.run(
		["pdftotext", str(pdf_path), "-"],
		check=True,
		capture_output=True,
		text=True,
	).stdout
	if len(text_content.strip()) < 100:
		raise RuntimeError(f"{pdf_path}: PDF does not contain enough selectable text")
	build_lib.syllabus_content.verify_required_section_titles(text_content, manifest, pdf_path)
	build_lib.syllabus_content.scan_text_for_secrets(text_content, str(pdf_path))
	accessibility_findings = []
	if re.search(r"^Tagged:\s+yes\s*$", metadata, re.MULTILINE | re.IGNORECASE) is None:
		accessibility_findings.append("PDF is not tagged")
	if not reader.outline:
		accessibility_findings.append("PDF has no heading bookmarks")
	for finding in accessibility_findings:
		print(f"Accessibility advisory: {pdf_path}: {finding}")
	return None


#============================================
def scan_docx_for_secrets(docx_path: pathlib.Path) -> None:
	"""Scan all XML text stored in one generated DOCX."""
	with zipfile.ZipFile(docx_path) as archive:
		xml_parts = []
		for member_name in archive.namelist():
			if member_name.endswith(".xml"):
				xml_parts.append(archive.read(member_name).decode("utf-8", errors="replace"))
	combined_xml = "\n".join(xml_parts)
	build_lib.syllabus_content.scan_text_for_secrets(combined_xml, str(docx_path))
	return None


#============================================
def check_tools() -> None:
	"""Require the lightweight document tools needed for complete exports."""
	for tool_name in ("pandoc", "pdfinfo", "pdftotext"):
		if shutil.which(tool_name) is None:
			raise RuntimeError(f"Required export tool is not installed: {tool_name}")
	if importlib.util.find_spec("weasyprint") is None:
		raise RuntimeError("Required PDF renderer is missing: install pip_requirements.txt")
	return None


#============================================
def publish_downloads(
	staged_downloads_dir: pathlib.Path,
	downloads_dir: pathlib.Path,
	expected_names: set[str],
) -> None:
	"""Publish one completely validated managed download set."""
	staged_paths = tuple(
		path
		for path in sorted(staged_downloads_dir.iterdir())
		if path.is_file() and path.suffix.lower() in MANAGED_DOWNLOAD_SUFFIXES
	)
	staged_names = {path.name for path in staged_paths}
	if staged_names != expected_names:
		raise RuntimeError(
			f"Staged downloads do not match the manifest set: {sorted(staged_names)}"
		)
	downloads_dir.mkdir(parents=True, exist_ok=True)
	# ASVS 2.3.3: validate the complete staged set before replacing published files.
	# ASVS 5.3.2: output names come only from validated manifest download basenames.
	for staged_path in staged_paths:
		final_path = downloads_dir / staged_path.name
		os.replace(staged_path, final_path)
	for artifact_path in downloads_dir.iterdir():
		if (
			artifact_path.is_file()
			and artifact_path.suffix.lower() in MANAGED_DOWNLOAD_SUFFIXES
			and artifact_path.name not in expected_names
		):
			artifact_path.unlink()
	return None


#============================================
def build_one_syllabus(
	manifest: build_lib.syllabus_model.SyllabusManifest,
	downloads_dir: pathlib.Path,
	reference_path: pathlib.Path,
	pdf_stylesheet_path: pathlib.Path,
	markdown_extensions: tuple[str, ...],
	markdown_extension_configs: dict[str, dict[object, object]],
	temporary_dir: pathlib.Path,
) -> tuple[pathlib.Path, pathlib.Path]:
	"""Build sibling DOCX and PDF outputs from one composed Markdown source."""
	combined_markdown = build_lib.syllabus_content.compose_markdown(manifest)
	docx_markdown_path = temporary_dir / f"{manifest.download_basename}_docx.md"
	docx_markdown = build_lib.syllabus_content.normalize_admonitions(combined_markdown)
	docx_markdown_path.write_text(docx_markdown, encoding="utf-8")
	docx_path = downloads_dir / f"{manifest.download_basename}.docx"
	run_pandoc_docx(docx_markdown_path, docx_path, reference_path, manifest)
	postprocess_docx(docx_path, manifest)
	verify_docx_output(docx_path, manifest)
	audit_docx_structure(docx_path, manifest)
	scan_docx_for_secrets(docx_path)
	html_path = temporary_dir / f"{manifest.download_basename}.html"
	run_markdown_html(
		combined_markdown,
		html_path,
		pdf_stylesheet_path,
		manifest,
		markdown_extensions,
		markdown_extension_configs,
	)
	pdf_path = downloads_dir / f"{manifest.download_basename}.pdf"
	run_weasyprint_pdf(html_path, pdf_path)
	verify_pdf_output(pdf_path, manifest)
	print(f"Built {docx_path.name} and {pdf_path.name}")
	outputs = (docx_path, pdf_path)
	return outputs
