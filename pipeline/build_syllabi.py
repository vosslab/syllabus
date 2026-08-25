"""Build complete DOCX and PDF syllabi from course manifests."""

# Standard Library
import os
import re
import sys
import html
import shutil
import pathlib
import zipfile
import argparse
import tempfile
import subprocess
import dataclasses
import importlib.util

# PIP3 modules
import docx
import yaml
import pypdf
import markdown
import docx.oxml
import docx.oxml.ns


SECRET_PATTERNS = (
	re.compile(r"zoom\.us/j/", re.IGNORECASE),
	re.compile(r"\bpwd=", re.IGNORECASE),
	re.compile(r"\b(?:passcode|password)\s*[:=]", re.IGNORECASE),
	re.compile(r"discord(?:\.gg|\.com/invite)/", re.IGNORECASE),
)
PROHIBITED_CONTROL_PATTERN = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
MARKDOWN_TABLE_SEPARATOR_PATTERN = re.compile(r":?-{3,}:?")
INCLUDE_LINE_PATTERN = re.compile(r'^--8<--[ \t]+"([^"\r\n]+)"[ \t]*$', re.MULTILINE)
SAFE_INCLUDE_PATH_PATTERN = re.compile(r"[A-Za-z0-9][A-Za-z0-9_./-]*\.md")
MANAGED_DOWNLOAD_SUFFIXES = {".docx", ".pdf"}
REQUIRED_LEARNING_TITLE = "# Learning Objectives, Outcomes, and Goals"
REQUIRED_LEARNING_MARKERS = (
	"## Roosevelt learning goals",
	"## Learning Objectives",
	"Students completing this course will have achieved:",
	"## Course Learning Outcomes",
	"Students completing this course will be able to:",
	"## Learning Goals",
	"Overall, this course aims to accomplish:",
)


@dataclasses.dataclass(frozen=True)
class SyllabusManifest:
	"""Validated paths and metadata for one complete syllabus."""

	path: pathlib.Path
	docs_root: pathlib.Path
	title: str
	course_code: str
	term: str
	author: str
	language: str
	download_basename: str
	sections: tuple[pathlib.Path, ...]
	shared_sections: tuple[pathlib.Path, ...]


#============================================
def get_repo_root() -> pathlib.Path:
	"""Return the repository root reported by Git."""
	completed = subprocess.run(
		["git", "rev-parse", "--show-toplevel"],
		check=True,
		capture_output=True,
		text=True,
	)
	repo_root = pathlib.Path(completed.stdout.strip())
	return repo_root


#============================================
def require_text(data: dict[object, object], key: str, manifest_path: pathlib.Path) -> str:
	"""Return one required, non-empty string field from manifest data."""
	value = data[key]
	if not isinstance(value, str) or not value.strip():
		raise ValueError(f"{manifest_path}: {key} must be a non-empty string")
	text_value = value.strip()
	return text_value


#============================================
def resolve_sources(
	data: dict[object, object],
	key: str,
	manifest_path: pathlib.Path,
	docs_root: pathlib.Path,
) -> tuple[pathlib.Path, ...]:
	"""Resolve and validate an ordered manifest source list."""
	value = data[key]
	if not isinstance(value, list) or not value:
		raise ValueError(f"{manifest_path}: {key} must be a non-empty list")
	resolved_paths = []
	for item in value:
		if not isinstance(item, str) or not item.strip():
			raise ValueError(f"{manifest_path}: {key} entries must be non-empty strings")
		candidate = (manifest_path.parent / item).resolve()
		if not candidate.is_relative_to(docs_root.resolve()):
			raise ValueError(f"{manifest_path}: source escapes site_docs: {item}")
		if not candidate.is_file():
			raise FileNotFoundError(f"{manifest_path}: missing source: {item}")
		resolved_paths.append(candidate)
	sources = tuple(resolved_paths)
	return sources


#============================================
def load_manifest(manifest_path: pathlib.Path, docs_root: pathlib.Path) -> SyllabusManifest:
	"""Load one YAML manifest and reject incomplete or unsafe values."""
	loaded = yaml.safe_load(manifest_path.read_text(encoding="utf-8"))
	if not isinstance(loaded, dict):
		raise ValueError(f"{manifest_path}: manifest root must be a mapping")
	title = require_text(loaded, "title", manifest_path)
	course_code = require_text(loaded, "course_code", manifest_path)
	term = require_text(loaded, "term", manifest_path)
	author = require_text(loaded, "author", manifest_path)
	language = require_text(loaded, "language", manifest_path)
	download_basename = require_text(loaded, "download_basename", manifest_path)
	if re.fullmatch(r"[A-Z0-9_]+", download_basename) is None:
		raise ValueError(f"{manifest_path}: download_basename must use A-Z, 0-9, and underscores")
	sections = resolve_sources(loaded, "sections", manifest_path, docs_root)
	shared_sections = resolve_sources(loaded, "shared_sections", manifest_path, docs_root)
	validate_course_learning_framework(sections, docs_root)
	manifest = SyllabusManifest(
		path=manifest_path,
		docs_root=docs_root.resolve(),
		title=title,
		course_code=course_code,
		term=term,
		author=author,
		language=language,
		download_basename=download_basename,
		sections=sections,
		shared_sections=shared_sections,
	)
	return manifest


#============================================
def scan_text_for_secrets(text_value: str, source_label: str) -> None:
	"""Reject public content that resembles meeting credentials or invite links."""
	for pattern in SECRET_PATTERNS:
		match = pattern.search(text_value)
		if match is not None:
			raise ValueError(f"{source_label}: prohibited public credential pattern: {match.group(0)}")
	return None


#============================================
def scan_text_for_prohibited_controls(text_value: str, source_label: str) -> None:
	"""Reject control characters that can change Markdown line structure."""
	match = PROHIBITED_CONTROL_PATTERN.search(text_value)
	if match is not None:
		codepoint = ord(match.group(0))
		raise ValueError(f"{source_label}: prohibited control character: U+{codepoint:04X}")
	return None


#============================================
def split_markdown_table_row(line: str) -> tuple[str, ...]:
	"""Return cells from one leading-and-trailing-pipe Markdown row."""
	stripped = line.strip()
	if not stripped.startswith("|") or not stripped.endswith("|"):
		raise ValueError("Markdown table rows must start and end with a pipe")
	cell_text = stripped[1:-1]
	cells = tuple(cell.strip() for cell in re.split(r"(?<!\\)\|", cell_text))
	return cells


#============================================
def count_markdown_tables(markdown_text: str) -> int:
	"""Count pipe tables by their header separator rows."""
	lines = markdown_text.split("\n")
	table_count = 0
	for line_index in range(1, len(lines)):
		current = lines[line_index].strip()
		previous = lines[line_index - 1].strip()
		if not current.startswith("|") or not current.endswith("|"):
			continue
		if not previous.startswith("|") or not previous.endswith("|"):
			continue
		separator_cells = split_markdown_table_row(current)
		if separator_cells and all(
			MARKDOWN_TABLE_SEPARATOR_PATTERN.fullmatch(cell) is not None
			for cell in separator_cells
		):
			table_count += 1
	return table_count


#============================================
def validate_markdown_tables(markdown_text: str, source_label: str) -> None:
	"""Require simple, rectangular Markdown tables with named header cells."""
	lines = markdown_text.split("\n")
	line_index = 0
	while line_index < len(lines):
		line = lines[line_index].strip()
		if not line.startswith("|") or not line.endswith("|"):
			line_index += 1
			continue
		block_start = line_index
		block_lines = []
		while line_index < len(lines):
			candidate = lines[line_index].strip()
			if not candidate.startswith("|") or not candidate.endswith("|"):
				break
			block_lines.append(lines[line_index])
			line_index += 1
		if len(block_lines) < 2:
			raise ValueError(f"{source_label}:{block_start + 1}: incomplete Markdown table")
		header_cells = split_markdown_table_row(block_lines[0])
		separator_cells = split_markdown_table_row(block_lines[1])
		if len(header_cells) < 2 or any(not cell for cell in header_cells):
			raise ValueError(f"{source_label}:{block_start + 1}: table headers must be named")
		if len(separator_cells) != len(header_cells) or any(
			MARKDOWN_TABLE_SEPARATOR_PATTERN.fullmatch(cell) is None
			for cell in separator_cells
		):
			raise ValueError(f"{source_label}:{block_start + 2}: invalid table separator row")
		for row_offset, row_line in enumerate(block_lines[2:], start=3):
			row_cells = split_markdown_table_row(row_line)
			if len(row_cells) != len(header_cells):
				raise ValueError(
					f"{source_label}:{block_start + row_offset}: inconsistent table columns"
				)
	return None


#============================================
def scan_public_sources(docs_root: pathlib.Path) -> None:
	"""Scan tracked-shape public text sources before generating downloads."""
	for suffix in ("*.md", "*.yml", "*.yaml"):
		for source_path in sorted(docs_root.rglob(suffix)):
			if "downloads" in source_path.parts:
				continue
			content = source_path.read_text(encoding="utf-8")
			scan_text_for_prohibited_controls(content, str(source_path))
			scan_text_for_secrets(content, str(source_path))
			if source_path.suffix == ".md":
				validate_markdown_tables(content, str(source_path))
	return None


#============================================
def require_public_only_repository(repo_root: pathlib.Path) -> None:
	"""Reject a private or ambiguous raw-content tree inside the repository."""
	raw_path = repo_root / "raw"
	if raw_path.exists():
		raise RuntimeError(
			f"{raw_path}: only public-safe canonical content belongs in this repository"
		)
	return None


#============================================
def require_single_content_authority(repo_root: pathlib.Path) -> None:
	"""Reject a parallel Markdown or manifest tree below templates."""
	templates_root = repo_root / "templates"
	for suffix in ("*.md", "*.yml", "*.yaml"):
		for source_path in sorted(templates_root.rglob(suffix)):
			# ASVS 2.1.1: enforce the documented site_docs-only source boundary.
			raise RuntimeError(
				f"{source_path}: live syllabus content belongs only under site_docs"
			)
	return None


#============================================
def expand_shared_includes(
	markdown_text: str,
	source_path: pathlib.Path,
	docs_root: pathlib.Path,
) -> str:
	"""Expand one level of restricted Markdown includes below site_docs."""
	resolved_docs_root = docs_root.resolve()
	resolved_source_path = source_path.resolve()
	if not resolved_source_path.is_relative_to(resolved_docs_root):
		raise ValueError(f"{source_path}: include source escapes site_docs")

	# ASVS 2.2.1 and 5.3.2: allow only simple Markdown paths below site_docs.
	def replace_include(match: re.Match[str]) -> str:
		include_name = match.group(1)
		include_parts = pathlib.PurePosixPath(include_name).parts
		if SAFE_INCLUDE_PATH_PATTERN.fullmatch(include_name) is None or ".." in include_parts:
			raise ValueError(f"{source_path}: unsafe Markdown include: {include_name}")
		include_path = (resolved_docs_root / include_name).resolve()
		if not include_path.is_relative_to(resolved_docs_root):
			raise ValueError(f"{source_path}: include escapes site_docs: {include_name}")
		if not include_path.is_file():
			raise FileNotFoundError(f"{source_path}: missing Markdown include: {include_name}")
		include_markdown = include_path.read_text(encoding="utf-8")
		if not include_markdown.strip():
			raise ValueError(f"{include_path}: Markdown include must not be empty")
		if INCLUDE_LINE_PATTERN.search(include_markdown) is not None:
			raise ValueError(f"{include_path}: nested Markdown includes are not supported")
		return include_markdown.strip()

	expanded = INCLUDE_LINE_PATTERN.sub(replace_include, markdown_text)
	return expanded


#============================================
def validate_course_learning_framework(
	sections: tuple[pathlib.Path, ...],
	docs_root: pathlib.Path,
) -> None:
	"""Require the four ordered learning sections and Roosevelt goal bullets."""
	framework_paths = tuple(
		path for path in sections if path.name == "COURSE_LEARNING_FRAMEWORK.md"
	)
	if len(framework_paths) != 1:
		raise ValueError("sections must contain exactly one COURSE_LEARNING_FRAMEWORK.md")
	framework_path = framework_paths[0]
	markdown = framework_path.read_text(encoding="utf-8")
	markdown = expand_shared_includes(markdown, framework_path, docs_root)
	if not markdown.startswith(REQUIRED_LEARNING_TITLE + "\n"):
		raise ValueError(
			f"{framework_path}: title must be {REQUIRED_LEARNING_TITLE.removeprefix('# ')}"
		)
	marker_positions = []
	for marker in REQUIRED_LEARNING_MARKERS:
		position = markdown.find(marker)
		if position < 0:
			raise ValueError(f"{framework_path}: missing required learning marker: {marker}")
		marker_positions.append(position)
	if marker_positions != sorted(marker_positions):
		raise ValueError(f"{framework_path}: required learning sections are out of order")
	roosevelt_start = marker_positions[0] + len(REQUIRED_LEARNING_MARKERS[0])
	roosevelt_end = marker_positions[1]
	roosevelt_markdown = markdown[roosevelt_start:roosevelt_end]
	if re.search(r"^[-*+]\s+\S", roosevelt_markdown, re.MULTILINE) is None:
		raise ValueError(f"{framework_path}: Roosevelt learning goals must be bullet points")
	return None


#============================================
def normalize_admonitions(markdown: str) -> str:
	"""Convert Material admonitions into portable Markdown for Pandoc."""
	lines = markdown.splitlines()
	normalized = []
	in_admonition = False
	for line in lines:
		match = re.fullmatch(r'!!!\s+[-\w]+(?:\s+"([^"]+)")?\s*', line)
		if match is not None:
			title = match.group(1)
			if title is None:
				title = "Note"
			normalized.append(f"**{title}**")
			in_admonition = True
			continue
		if in_admonition and line.startswith("    "):
			normalized.append(line[4:])
			continue
		if in_admonition and not line.strip():
			normalized.append("")
			continue
		in_admonition = False
		normalized.append(line)
	result = "\n".join(normalized)
	return result


#============================================
def remove_heading_sections(markdown_text: str, heading_names: tuple[str, ...]) -> str:
	"""Remove level-two web-navigation sections from composed documents."""
	heading_options = "|".join(re.escape(name) for name in heading_names)
	pattern = rf"^## (?:{heading_options})\s*$.*?(?=^## |\Z)"
	without_sections = re.sub(pattern, "", markdown_text, flags=re.MULTILINE | re.DOTALL)
	return without_sections


#============================================
def rewrite_document_links(
	markdown_text: str,
	source_path: pathlib.Path,
	document_anchors: dict[pathlib.Path, str],
) -> str:
	"""Point links between included Markdown files at complete-document sections."""
	def replace_link(match: re.Match[str]) -> str:
		"""Rewrite one link whose target is represented in the document."""
		relative_target = match.group(1)
		fragment = match.group(2)
		resolved_target = (source_path.parent / relative_target).resolve()
		# ASVS 5.3.2: rewrite only paths already validated into this manifest.
		if resolved_target not in document_anchors:
			return match.group(0)
		if fragment:
			document_target = fragment
		else:
			document_target = f"#{document_anchors[resolved_target]}"
		replacement = f"({document_target})"
		return replacement

	pattern = r"\(([^()\s]+\.md)(#[^()\s]+)?\)"
	rewritten = re.sub(pattern, replace_link, markdown_text)
	return rewritten


#============================================
def prepare_section(markdown: str, is_overview: bool, anchor: str) -> str:
	"""Remove web-only controls and demote headings for the merged document."""
	without_downloads = re.sub(
		r'<div class="syllabus-downloads".*?</div>\s*',
		"",
		markdown,
		flags=re.DOTALL,
	)
	if is_overview:
		without_navigation = remove_heading_sections(
			without_downloads,
			("Course pages", "Find what you need"),
		)
	else:
		without_navigation = without_downloads
	lines = []
	anchored_heading = False
	for line in without_navigation.splitlines():
		match = re.match(r"^(#{1,6})(\s+.*)$", line)
		if match is None:
			lines.append(line)
			continue
		level = min(len(match.group(1)) + 1, 6)
		heading_suffix = match.group(2)
		if not anchored_heading:
			heading_suffix += f" {{#{anchor}}}"
			anchored_heading = True
		lines.append("#" * level + heading_suffix)
	prepared = "\n".join(lines).strip() + "\n"
	return prepared


#============================================
def load_markdown_configuration(
	config_path: pathlib.Path,
) -> tuple[tuple[str, ...], dict[str, dict[object, object]]]:
	"""Load the Python-Markdown extension stack used by the MkDocs site."""
	loaded = yaml.safe_load(config_path.read_text(encoding="utf-8"))
	if not isinstance(loaded, dict):
		raise ValueError(f"{config_path}: configuration root must be a mapping")
	if "markdown_extensions" not in loaded:
		raise ValueError(f"{config_path}: missing markdown_extensions")
	configured_extensions = loaded["markdown_extensions"]
	if not isinstance(configured_extensions, list) or not configured_extensions:
		raise ValueError(f"{config_path}: markdown_extensions must be a non-empty list")
	extension_names = []
	extension_configs = {}
	for item in configured_extensions:
		if isinstance(item, str):
			extension_names.append(item)
			continue
		if not isinstance(item, dict) or len(item) != 1:
			raise ValueError(f"{config_path}: invalid markdown extension entry: {item!r}")
		extension_name, settings = next(iter(item.items()))
		if not isinstance(extension_name, str):
			raise ValueError(f"{config_path}: markdown extension names must be strings")
		if settings is None:
			settings = {}
		if not isinstance(settings, dict):
			raise ValueError(f"{config_path}: {extension_name} settings must be a mapping")
		extension_names.append(extension_name)
		extension_configs[extension_name] = dict(settings)
	extensions = tuple(extension_names)
	return extensions, extension_configs


#============================================
def get_section_title(markdown: str, source_path: pathlib.Path) -> str:
	"""Return the first level-one heading from one section source."""
	match = re.search(r"^#\s+(.+?)\s*$", markdown, re.MULTILINE)
	if match is None:
		raise ValueError(f"{source_path}: section must begin with a level-one heading")
	title = match.group(1)
	return title


#============================================
def verify_required_section_titles(
	output_text: str,
	manifest: SyllabusManifest,
	output_path: pathlib.Path,
) -> None:
	"""Require every manifest source heading in one generated document."""
	normalized_output = re.sub(r"\s+", " ", output_text).casefold()
	missing_titles = []
	for source_path in manifest.sections + manifest.shared_sections:
		source_markdown = source_path.read_text(encoding="utf-8")
		section_title = get_section_title(source_markdown, source_path)
		normalized_title = re.sub(r"\s+", " ", section_title).casefold()
		if normalized_title not in normalized_output:
			missing_titles.append(section_title)
	if missing_titles:
		missing_text = ", ".join(missing_titles)
		raise RuntimeError(f"{output_path}: missing manifest sections: {missing_text}")
	return None


#============================================
def verify_download_links(manifest: SyllabusManifest, downloads_dir: pathlib.Path) -> None:
	"""Require overview download targets derived from manifest output names."""
	overview_path = manifest.sections[0]
	overview_markdown = overview_path.read_text(encoding="utf-8")
	for suffix in (".pdf", ".docx"):
		target_path = downloads_dir / f"{manifest.download_basename}{suffix}"
		relative_path = pathlib.Path(os.path.relpath(target_path, overview_path.parent)).as_posix()
		if relative_path not in overview_markdown:
			raise RuntimeError(f"{overview_path}: missing complete-download link: {relative_path}")
	return None


#============================================
def compose_markdown(manifest: SyllabusManifest) -> str:
	"""Compose course and shared sources in manifest order."""
	parts = []
	contents = ["# Contents", ""]
	document_anchors = {}
	for index, section_path in enumerate(manifest.sections):
		anchor = "course-overview" if index == 0 else section_path.stem.lower().replace("_", "-")
		document_anchors[section_path.resolve()] = anchor
	for section_path in manifest.shared_sections:
		anchor = section_path.stem.lower().replace("_", "-")
		document_anchors[section_path.resolve()] = anchor
	# The instructor policy route is embedded under this heading in COURSE_DETAILS.md.
	instructor_route_path = manifest.path.parent.parent / "policies" / "INSTRUCTOR_INFORMATION.md"
	if instructor_route_path.is_file():
		document_anchors[instructor_route_path.resolve()] = "instructor-information"
	for index, section_path in enumerate(manifest.sections):
		markdown = section_path.read_text(encoding="utf-8")
		markdown = expand_shared_includes(markdown, section_path, manifest.docs_root)
		markdown = rewrite_document_links(markdown, section_path, document_anchors)
		anchor = document_anchors[section_path.resolve()]
		title = "Course overview" if index == 0 else get_section_title(markdown, section_path)
		contents.append(f"- [{title}](#{anchor})")
		parts.append(prepare_section(markdown, is_overview=index == 0, anchor=anchor))
	for section_path in manifest.shared_sections:
		markdown = section_path.read_text(encoding="utf-8")
		markdown = expand_shared_includes(markdown, section_path, manifest.docs_root)
		markdown = rewrite_document_links(markdown, section_path, document_anchors)
		if section_path.name == "POLICIES.md":
			markdown = remove_heading_sections(markdown, ("Policy topics", "Student support"))
		anchor = document_anchors[section_path.resolve()]
		title = get_section_title(markdown, section_path)
		contents.append(f"- [{title}](#{anchor})")
		parts.append(prepare_section(markdown, is_overview=False, anchor=anchor))
	contents.append("")
	combined = "\n".join(contents) + "\n" + "\n\n".join(parts)
	scan_text_for_secrets(combined, str(manifest.path))
	return combined


#============================================
def mark_table_header(row: object) -> None:
	"""Mark one DOCX row as a repeating semantic table header."""
	properties = row._tr.get_or_add_trPr()
	if properties.find(docx.oxml.ns.qn("w:tblHeader")) is None:
		header = docx.oxml.OxmlElement("w:tblHeader")
		header.set(docx.oxml.ns.qn("w:val"), "true")
		properties.append(header)
	return None


#============================================
def format_table_header(row: object) -> None:
	"""Make table headers distinct through both bold text and light shading."""
	mark_table_header(row)
	for cell in row.cells:
		cell_properties = cell._tc.get_or_add_tcPr()
		shading = cell_properties.find(docx.oxml.ns.qn("w:shd"))
		if shading is None:
			shading = docx.oxml.OxmlElement("w:shd")
			cell_properties.append(shading)
		shading.set(docx.oxml.ns.qn("w:fill"), "E6E6E6")
		for paragraph in cell.paragraphs:
			for run in paragraph.runs:
				run.bold = True
	return None


#============================================
def prevent_row_split(row: object) -> None:
	"""Keep a DOCX table row from splitting across pages."""
	properties = row._tr.get_or_add_trPr()
	if properties.find(docx.oxml.ns.qn("w:cantSplit")) is None:
		cant_split = docx.oxml.OxmlElement("w:cantSplit")
		properties.append(cant_split)
	return None


#============================================
def set_document_language(document: object, language_code: str) -> None:
	"""Set the proofing language on paragraph and character styles."""
	for style in document.styles:
		run_properties = style.element.get_or_add_rPr()
		language = run_properties.find(docx.oxml.ns.qn("w:lang"))
		if language is None:
			language = docx.oxml.OxmlElement("w:lang")
			run_properties.append(language)
		language.set(docx.oxml.ns.qn("w:val"), language_code)
	return None


#============================================
def postprocess_docx(docx_path: pathlib.Path, manifest: SyllabusManifest) -> None:
	"""Apply metadata, language, and table accessibility properties."""
	document = docx.Document(docx_path)
	document.core_properties.title = f"{manifest.course_code}: {manifest.title} - {manifest.term}"
	document.core_properties.author = manifest.author
	document.core_properties.subject = "Complete course syllabus"
	document.core_properties.language = manifest.language
	set_document_language(document, manifest.language)
	for table in document.tables:
		table.style = "Table"
		table.autofit = True
		if table.rows:
			format_table_header(table.rows[0])
		for row in table.rows:
			prevent_row_split(row)
	document.save(docx_path)
	return None


#============================================
def audit_docx_structure(docx_path: pathlib.Path, manifest: SyllabusManifest) -> None:
	"""Report non-blocking DOCX accessibility findings."""
	document = docx.Document(docx_path)
	accessibility_findings = []
	expected_title = f"{manifest.course_code}: {manifest.title} - {manifest.term}"
	if document.core_properties.title != expected_title:
		accessibility_findings.append("missing expected document title metadata")
	if document.core_properties.language != manifest.language:
		accessibility_findings.append("missing expected document language metadata")
	heading_paragraphs = [
		paragraph
		for paragraph in document.paragraphs
		if paragraph.style.name.startswith("Heading ")
	]
	if not heading_paragraphs:
		accessibility_findings.append("no semantic heading styles found")
	if not document.tables:
		accessibility_findings.append("no semantic tables found")
	for table in document.tables:
		header_properties = table.rows[0]._tr.get_or_add_trPr()
		if header_properties.find(docx.oxml.ns.qn("w:tblHeader")) is None:
			accessibility_findings.append("table is missing a repeating header row")
	for finding in accessibility_findings:
		print(f"Accessibility advisory: {docx_path}: {finding}")
	return None


#============================================
def verify_docx_output(docx_path: pathlib.Path, manifest: SyllabusManifest) -> None:
	"""Require a readable DOCX package containing every manifest section."""
	document = docx.Document(docx_path)
	paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
	verify_required_section_titles(paragraph_text, manifest, docx_path)
	combined_markdown = compose_markdown(manifest)
	validate_markdown_tables(combined_markdown, str(manifest.path))
	expected_table_count = count_markdown_tables(combined_markdown)
	if len(document.tables) != expected_table_count:
		raise RuntimeError(
			f"{docx_path}: expected {expected_table_count} tables, found {len(document.tables)}"
		)
	return None


#============================================
def run_pandoc_docx(
	markdown_path: pathlib.Path,
	docx_path: pathlib.Path,
	reference_path: pathlib.Path,
	manifest: SyllabusManifest,
) -> None:
	"""Generate one complete DOCX through Pandoc."""
	command = [
		"pandoc",
		str(markdown_path),
		"--from=markdown+pipe_tables",
		"--to=docx",
		f"--reference-doc={reference_path}",
		f"--metadata=title:{manifest.course_code}: {manifest.title}",
		f"--metadata=subtitle:{manifest.term}",
		f"--metadata=author:{manifest.author}",
		f"--metadata=lang:{manifest.language}",
		f"--output={docx_path}",
	]
	subprocess.run(command, check=True)
	return None


#============================================
def run_markdown_html(
	markdown_text: str,
	html_path: pathlib.Path,
	stylesheet_path: pathlib.Path,
	manifest: SyllabusManifest,
	extensions: tuple[str, ...],
	extension_configs: dict[str, dict[object, object]],
) -> None:
	"""Render semantic standalone HTML with the site's Markdown extension stack."""
	validate_markdown_tables(markdown_text, str(manifest.path))
	converter = markdown.Markdown(
		extensions=list(extensions),
		extension_configs=extension_configs,
		output_format="html5",
	)
	body_html = converter.convert(markdown_text)
	expected_table_count = count_markdown_tables(markdown_text)
	rendered_table_count = body_html.count("<table>")
	if rendered_table_count != expected_table_count:
		raise RuntimeError(
			f"{html_path}: expected {expected_table_count} tables, found {rendered_table_count}"
		)
	document_title = f"{manifest.course_code}: {manifest.title} - {manifest.term}"
	escaped_title = html.escape(document_title)
	escaped_course_title = html.escape(f"{manifest.course_code}: {manifest.title}")
	escaped_term = html.escape(manifest.term)
	escaped_author = html.escape(manifest.author)
	escaped_language = html.escape(manifest.language, quote=True)
	document = "\n".join(
		[
			"<!doctype html>",
			f'<html lang="{escaped_language}">',
			"<head>",
			'<meta charset="utf-8">',
			f"<title>{escaped_title}</title>",
			f'<meta name="author" content="{escaped_author}">',
			f'<link rel="stylesheet" href="{stylesheet_path.as_uri()}">',
			"</head>",
			'<body class="syllabus-document">',
			'<header id="title-block-header">',
			f"<h1>{escaped_course_title}</h1>",
			f'<p class="subtitle">{escaped_term}</p>',
			f'<p class="author">{escaped_author}</p>',
			"</header>",
			'<main id="syllabus-content">',
			body_html,
			"</main>",
			"</body>",
			"</html>",
			"",
		]
	)
	html_path.write_text(document, encoding="utf-8")
	return None


#============================================
def run_weasyprint_pdf(
	html_path: pathlib.Path,
	pdf_path: pathlib.Path,
) -> None:
	"""Render standalone syllabus HTML directly to a tagged PDF."""
	command = [
		sys.executable,
		"-m",
		"weasyprint",
		"--pdf-tags",
		str(html_path),
		str(pdf_path),
	]
	subprocess.run(command, check=True)
	if not pdf_path.is_file():
		raise RuntimeError(f"WeasyPrint did not create {pdf_path}")
	return None


#============================================
def verify_pdf_output(pdf_path: pathlib.Path, manifest: SyllabusManifest) -> None:
	"""Require a usable PDF and report non-blocking accessibility findings."""
	metadata = subprocess.run(
		["pdfinfo", str(pdf_path)],
		check=True,
		capture_output=True,
		text=True,
	).stdout
	if re.search(r"^Page size:\s+612 x 792 pts", metadata, re.MULTILINE) is None:
		raise RuntimeError(f"{pdf_path}: PDF is not US letter size")
	reader = pypdf.PdfReader(pdf_path)
	if not reader.pages:
		raise RuntimeError(f"{pdf_path}: PDF contains no pages")
	text_content = subprocess.run(
		["pdftotext", str(pdf_path), "-"],
		check=True,
		capture_output=True,
		text=True,
	).stdout
	if len(text_content.strip()) < 100:
		raise RuntimeError(f"{pdf_path}: PDF does not contain enough selectable text")
	verify_required_section_titles(text_content, manifest, pdf_path)
	scan_text_for_secrets(text_content, str(pdf_path))
	accessibility_findings = []
	if re.search(r"^Tagged:\s+yes\s*$", metadata, re.MULTILINE | re.IGNORECASE) is None:
		accessibility_findings.append("PDF is not tagged")
	if not reader.outline:
		accessibility_findings.append("PDF has no heading bookmarks")
	for finding in accessibility_findings:
		print(f"Accessibility advisory: {pdf_path}: {finding}")
	return None


#============================================
def scan_docx_for_secrets(docx_path: pathlib.Path) -> None:
	"""Scan all XML text stored in one generated DOCX."""
	with zipfile.ZipFile(docx_path) as archive:
		xml_parts = []
		for member_name in archive.namelist():
			if member_name.endswith(".xml"):
				xml_parts.append(archive.read(member_name).decode("utf-8", errors="replace"))
	combined_xml = "\n".join(xml_parts)
	scan_text_for_secrets(combined_xml, str(docx_path))
	return None


#============================================
def check_tools() -> None:
	"""Require the lightweight document tools needed for complete exports."""
	for tool_name in ("pandoc", "pdfinfo", "pdftotext"):
		if shutil.which(tool_name) is None:
			raise RuntimeError(f"Required export tool is not installed: {tool_name}")
	if importlib.util.find_spec("weasyprint") is None:
		raise RuntimeError("Required PDF renderer is missing: install pip_requirements.txt")
	return None


#============================================
def publish_downloads(
	staged_downloads_dir: pathlib.Path,
	downloads_dir: pathlib.Path,
	expected_names: set[str],
) -> None:
	"""Publish one completely validated managed download set."""
	staged_paths = tuple(
		path
		for path in sorted(staged_downloads_dir.iterdir())
		if path.is_file() and path.suffix.lower() in MANAGED_DOWNLOAD_SUFFIXES
	)
	staged_names = {path.name for path in staged_paths}
	if staged_names != expected_names:
		raise RuntimeError(
			f"Staged downloads do not match the manifest set: {sorted(staged_names)}"
		)
	downloads_dir.mkdir(parents=True, exist_ok=True)
	# ASVS 2.3.3: validate the complete staged set before replacing published files.
	# ASVS 5.3.2: output names come only from validated manifest download basenames.
	for staged_path in staged_paths:
		final_path = downloads_dir / staged_path.name
		os.replace(staged_path, final_path)
	for artifact_path in downloads_dir.iterdir():
		if (
			artifact_path.is_file()
			and artifact_path.suffix.lower() in MANAGED_DOWNLOAD_SUFFIXES
			and artifact_path.name not in expected_names
		):
			artifact_path.unlink()
	return None


#============================================
def build_one_syllabus(
	manifest: SyllabusManifest,
	downloads_dir: pathlib.Path,
	reference_path: pathlib.Path,
	pdf_stylesheet_path: pathlib.Path,
	markdown_extensions: tuple[str, ...],
	markdown_extension_configs: dict[str, dict[object, object]],
	temporary_dir: pathlib.Path,
) -> tuple[pathlib.Path, pathlib.Path]:
	"""Build sibling DOCX and PDF outputs from one composed Markdown source."""
	combined_markdown = compose_markdown(manifest)
	docx_markdown_path = temporary_dir / f"{manifest.download_basename}_docx.md"
	docx_markdown_path.write_text(normalize_admonitions(combined_markdown), encoding="utf-8")
	docx_path = downloads_dir / f"{manifest.download_basename}.docx"
	run_pandoc_docx(docx_markdown_path, docx_path, reference_path, manifest)
	postprocess_docx(docx_path, manifest)
	verify_docx_output(docx_path, manifest)
	audit_docx_structure(docx_path, manifest)
	scan_docx_for_secrets(docx_path)
	html_path = temporary_dir / f"{manifest.download_basename}.html"
	run_markdown_html(
		combined_markdown,
		html_path,
		pdf_stylesheet_path,
		manifest,
		markdown_extensions,
		markdown_extension_configs,
	)
	pdf_path = downloads_dir / f"{manifest.download_basename}.pdf"
	run_weasyprint_pdf(html_path, pdf_path)
	verify_pdf_output(pdf_path, manifest)
	print(f"Built {docx_path.name} and {pdf_path.name}")
	outputs = (docx_path, pdf_path)
	return outputs


#============================================
def archive_outputs(
	outputs_by_term: dict[str, list[pathlib.Path]],
	archive_dir: pathlib.Path,
) -> None:
	"""Create one archival ZIP file per term."""
	archive_dir.mkdir(parents=True, exist_ok=True)
	for term, output_paths in sorted(outputs_by_term.items()):
		term_slug = re.sub(r"[^A-Z0-9]+", "_", term.upper()).strip("_")
		archive_path = archive_dir / f"{term_slug}_SYLLABI.zip"
		with zipfile.ZipFile(archive_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
			for output_path in sorted(output_paths):
				archive.write(output_path, arcname=output_path.name)
		print(f"Archived {archive_path}")
	return None


#============================================
def parse_args() -> argparse.Namespace:
	"""Parse the optional archive switch."""
	parser = argparse.ArgumentParser(description=__doc__)
	parser.add_argument(
		"--archive",
		action="store_true",
		help="also create one ZIP archive per term under output/archive",
	)
	args = parser.parse_args()
	return args


#============================================
def main() -> None:
	"""Build all course manifests into complete downloadable files."""
	args = parse_args()
	repo_root = get_repo_root()
	docs_root = repo_root / "site_docs"
	downloads_dir = docs_root / "downloads"
	reference_path = repo_root / "pipeline" / "syllabus_reference.docx"
	pdf_stylesheet_path = docs_root / "assets" / "stylesheets" / "syllabus_pdf.css"
	mkdocs_config_path = repo_root / "mkdocs.yml"
	if not reference_path.is_file():
		raise FileNotFoundError(
			f"Missing {reference_path}. Run pipeline/create_syllabus_reference_docx.py first."
		)
	if not pdf_stylesheet_path.is_file():
		raise FileNotFoundError(f"Missing PDF stylesheet: {pdf_stylesheet_path}")
	if not mkdocs_config_path.is_file():
		raise FileNotFoundError(f"Missing MkDocs configuration: {mkdocs_config_path}")
	check_tools()
	require_public_only_repository(repo_root)
	require_single_content_authority(repo_root)
	scan_public_sources(docs_root)
	markdown_extensions, markdown_extension_configs = load_markdown_configuration(
		mkdocs_config_path
	)
	manifest_paths = sorted(docs_root.rglob("syllabus.yml"))
	if not manifest_paths:
		raise RuntimeError("No syllabus.yml manifests found under site_docs")
	manifests = [load_manifest(manifest_path, docs_root) for manifest_path in manifest_paths]
	for manifest in manifests:
		verify_download_links(manifest, downloads_dir)
	expected_names = {
		f"{manifest.download_basename}{suffix}"
		for manifest in manifests
		for suffix in MANAGED_DOWNLOAD_SUFFIXES
	}
	outputs_by_term: dict[str, list[pathlib.Path]] = {}
	with tempfile.TemporaryDirectory(
		prefix=".syllabus_build_",
		dir=docs_root,
	) as temporary_name:
		temporary_dir = pathlib.Path(temporary_name)
		staged_downloads_dir = temporary_dir / "downloads"
		staged_downloads_dir.mkdir()
		for manifest in manifests:
			outputs = build_one_syllabus(
				manifest,
				staged_downloads_dir,
				reference_path,
				pdf_stylesheet_path,
				markdown_extensions,
				markdown_extension_configs,
				temporary_dir,
			)
			if manifest.term not in outputs_by_term:
				outputs_by_term[manifest.term] = []
			final_outputs = [downloads_dir / output.name for output in outputs]
			outputs_by_term[manifest.term].extend(final_outputs)
		publish_downloads(staged_downloads_dir, downloads_dir, expected_names)
	if args.archive:
		archive_outputs(outputs_by_term, repo_root / "output" / "archive")
	return None


if __name__ == "__main__":
	main()
