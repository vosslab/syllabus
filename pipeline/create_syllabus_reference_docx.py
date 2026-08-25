"""Create the clean reference DOCX used by the syllabus export pipeline."""

# Standard Library
import pathlib
import subprocess

# PIP3 modules
import docx
import docx.enum.style
import docx.enum.text
import docx.oxml
import docx.oxml.ns
import docx.shared


#============================================
def get_repo_root() -> pathlib.Path:
	"""Return the repository root reported by Git."""
	completed = subprocess.run(
		["git", "rev-parse", "--show-toplevel"],
		check=True,
		capture_output=True,
		text=True,
	)
	repo_root = pathlib.Path(completed.stdout.strip())
	return repo_root


#============================================
def set_font(style: object, name: str, size_points: float, bold: bool = False) -> None:
	"""Set a style's readable Latin font properties."""
	style.font.name = name
	style.font.size = docx.shared.Pt(size_points)
	style.font.bold = bold
	run_properties = style.element.get_or_add_rPr()
	run_fonts = run_properties.get_or_add_rFonts()
	run_fonts.set(docx.oxml.ns.qn("w:ascii"), name)
	run_fonts.set(docx.oxml.ns.qn("w:hAnsi"), name)
	return None


#============================================
def add_page_number(paragraph: object) -> None:
	"""Add a centered PAGE field to a footer paragraph."""
	paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
	run = paragraph.add_run("Page ")
	field_begin = docx.oxml.OxmlElement("w:fldChar")
	field_begin.set(docx.oxml.ns.qn("w:fldCharType"), "begin")
	instruction = docx.oxml.OxmlElement("w:instrText")
	instruction.set(docx.oxml.ns.qn("xml:space"), "preserve")
	instruction.text = " PAGE "
	field_end = docx.oxml.OxmlElement("w:fldChar")
	field_end.set(docx.oxml.ns.qn("w:fldCharType"), "end")
	run._r.append(field_begin)
	run._r.append(instruction)
	run._r.append(field_end)
	return None


#============================================
def configure_styles(document: object) -> None:
	"""Configure body, heading, title, and table styles for long syllabi."""
	normal = document.styles["Normal"]
	set_font(normal, "Arial", 10.5)
	normal.paragraph_format.space_after = docx.shared.Pt(6)
	normal.paragraph_format.line_spacing = 1.08
	normal.paragraph_format.widow_control = True
	compact = document.styles["Compact"]
	compact.paragraph_format.space_after = docx.shared.Pt(0)
	compact.paragraph_format.line_spacing = 1.0

	title = document.styles["Title"]
	set_font(title, "Arial", 22, bold=True)
	title.font.color.rgb = docx.shared.RGBColor(0x17, 0x36, 0x5D)
	title.paragraph_format.space_after = docx.shared.Pt(18)
	subtitle = document.styles["Subtitle"]
	set_font(subtitle, "Arial", 12)
	subtitle.font.color.rgb = docx.shared.RGBColor(0x00, 0x00, 0x00)
	subtitle.font.italic = False
	hyperlink = document.styles["Hyperlink"]
	hyperlink.font.color.rgb = docx.shared.RGBColor(0x00, 0x4A, 0x83)
	hyperlink.font.underline = True

	heading_sizes = {1: 17, 2: 14, 3: 12, 4: 10.5}
	for level, size in heading_sizes.items():
		heading = document.styles[f"Heading {level}"]
		set_font(heading, "Arial", size, bold=True)
		heading.font.color.rgb = docx.shared.RGBColor(0x00, 0x00, 0x00)
		heading.paragraph_format.keep_with_next = True
		heading.paragraph_format.space_before = docx.shared.Pt(12)
		heading.paragraph_format.space_after = docx.shared.Pt(4)

	for style in document.styles:
		if style.type not in {
			docx.enum.style.WD_STYLE_TYPE.CHARACTER,
			docx.enum.style.WD_STYLE_TYPE.PARAGRAPH,
		}:
			continue
		run_properties = style.element.get_or_add_rPr()
		language = run_properties.find(docx.oxml.ns.qn("w:lang"))
		if language is None:
			language = docx.oxml.OxmlElement("w:lang")
			run_properties.append(language)
		language.set(docx.oxml.ns.qn("w:val"), "en-US")
	return None


#============================================
def configure_page(document: object) -> None:
	"""Set letter-sized pages, practical margins, and page-number footers."""
	section = document.sections[0]
	section.page_width = docx.shared.Inches(8.5)
	section.page_height = docx.shared.Inches(11)
	section.top_margin = docx.shared.Inches(0.75)
	section.bottom_margin = docx.shared.Inches(0.75)
	section.left_margin = docx.shared.Inches(0.8)
	section.right_margin = docx.shared.Inches(0.8)
	add_page_number(section.footer.paragraphs[0])
	return None


#============================================
def create_reference_docx(output_path: pathlib.Path) -> None:
	"""Create a reference DOCX with accessible, restrained document styles."""
	output_path.parent.mkdir(parents=True, exist_ok=True)
	subprocess.run(
		[
			"pandoc",
			"-o",
			str(output_path),
			"--print-default-data-file",
			"reference.docx",
		],
		check=True,
	)
	document = docx.Document(output_path)
	configure_page(document)
	configure_styles(document)
	document.core_properties.title = "Complete course syllabus"
	document.core_properties.author = "Neil R. Voss"
	document.core_properties.subject = "Course syllabus"
	document.core_properties.language = "en-US"
	document.save(output_path)
	return None


#============================================
def main() -> None:
	"""Write the tracked reference document beside the export pipeline."""
	repo_root = get_repo_root()
	output_path = repo_root / "pipeline" / "syllabus_reference.docx"
	create_reference_docx(output_path)
	print(f"Created {output_path.relative_to(repo_root)}")
	return None


if __name__ == "__main__":
	main()
