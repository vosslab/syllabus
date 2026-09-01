"""Focused unit tests for syllabus composition and publication safety."""

# Standard Library
import pathlib

# PIP3 modules
import docx
import docx.oxml.ns
import pytest

# local repo modules
import build_lib.syllabus_model
import build_lib.syllabus_content
import build_lib.syllabus_rendering


#============================================
def test_course_document_basename_follows_public_filename_contract() -> None:
	"""Semantic course metadata produces the exact student-facing filename convention."""
	basename = build_lib.syllabus_model.format_course_document_basename(
		"BIOL 351/451",
		"Fall 2026",
		"Syllabus",
	)
	assert basename == "Voss-BIOL_351_451-Fall_2026-Syllabus"


#============================================
def test_assessment_examples_url_accepts_only_official_subject_routes() -> None:
	"""Assessment examples remain on the official OER subject routes."""
	manifest_path = pathlib.Path("syllabus.yml")
	valid_url = "https://biologyproblems.org/biochemistry/"
	assert (
		build_lib.syllabus_model.validate_assessment_examples_url(valid_url, manifest_path)
		== valid_url
	)
	with pytest.raises(ValueError, match="assessment_examples_url"):
		build_lib.syllabus_model.validate_assessment_examples_url(
			"javascript:alert(1)", manifest_path
		)


#============================================
def test_course_point_plan_derives_total_and_approximate_shares() -> None:
	"""One point edit drives every displayed arithmetic value."""
	manifest_path = pathlib.Path("syllabus.yml")
	point_plan = build_lib.syllabus_model.resolve_course_point_plan(
		{
			"course_point_plan": [
				{"assessment": "Research & analysis", "points": 2},
				{"assessment": "Final exam", "points": 1},
			]
		},
		manifest_path,
	)
	rendered = build_lib.syllabus_content.render_course_point_plan(point_plan)
	expected = (
		"| Assessment | Possible points | Approximate share | Your points |\n"
		"| --- | ---: | ---: | ---: |\n"
		"| Research &amp; analysis | 2 | 66.7% | |\n"
		"| Final exam | 1 | 33.3% | |\n"
		"| **Total** | **3** | **100%** | |"
	)
	assert rendered == expected


#============================================
def test_course_point_plan_requires_one_coursework_marker(tmp_path: pathlib.Path) -> None:
	"""Configured point data cannot be silently omitted from coursework."""
	coursework_path = tmp_path / "ASSIGNMENTS_AND_GRADING.md"
	manifest = build_lib.syllabus_model.SyllabusManifest(
		path=tmp_path / "syllabus.yml",
		docs_root=tmp_path,
		title="Course title",
		short_name="Course",
		course_code="BIOL 000",
		term="Fall 2099",
		author="Instructor",
		language="en-US",
		course_color="#007849",
		sections=(coursework_path,),
		shared_sections=(),
		lab_status="no_lab",
		course_point_plan=(
			build_lib.syllabus_model.CoursePointPlanEntry("Exam", 100),
		),
	)
	with pytest.raises(ValueError, match="expected exactly one course point-plan marker"):
		build_lib.syllabus_content.apply_course_point_plan(
			"# Coursework and grades\n",
			coursework_path,
			manifest,
		)


#============================================
def test_prepare_section_removes_web_only_content() -> None:
	"""The archival source excludes web controls while retaining extension syntax."""
	markdown = (
		"# Course title\n\n"
		'<div class="syllabus-downloads"><a href="file.pdf">Download PDF</a></div>\n\n'
		'!!! warning "Review required"\n\n    Confirm the grading plan.\n\n'
		"## Course pages\n\n- [Details](COURSE_DETAILS.md)\n\n"
		"## Course at a glance\n\n| Course | Term |\n| --- | --- |\n| Test | Fall |\n"
	)
	prepared = build_lib.syllabus_content.prepare_section(markdown, True, "course-overview")
	expected = (
		"## Course title {#course-overview}\n\n"
		'!!! warning "Review required"\n\n'
		"    Confirm the grading plan.\n\n"
		"### Course at a glance\n\n"
		"| Course | Term |\n| --- | --- |\n| Test | Fall |\n"
	)
	assert prepared == expected


#============================================
def test_download_links_cover_course_and_term_pages(tmp_path: pathlib.Path) -> None:
	"""Each manifest output is linked from both student landing pages."""
	docs_root = tmp_path / "site_docs"
	term_path = docs_root / "fall_20xx"
	course_path = term_path / "course"
	fragments_path = term_path / "shared" / "fragments"
	downloads_dir = docs_root / "downloads"
	course_path.mkdir(parents=True)
	fragments_path.mkdir(parents=True)
	overview_path = course_path / "index.md"
	term_overview_path = term_path / "index.md"
	term_courses_path = fragments_path / "TERM_COURSES.md"
	overview_path.write_text(
		"../../downloads/Voss-BIOL_000-Fall_2099-Syllabus.pdf\n"
		"../../downloads/Voss-BIOL_000-Fall_2099-Syllabus.docx\n",
		encoding="utf-8",
	)
	term_overview_path.write_text(
		'--8<-- "fall_20xx/shared/fragments/TERM_COURSES.md"\n',
		encoding="utf-8",
	)
	term_courses_path.write_text(
		'<a href="../../../downloads/Voss-BIOL_000-Fall_2099-Syllabus.pdf">PDF</a>\n'
		'<a href="../../../downloads/Voss-BIOL_000-Fall_2099-Syllabus.docx">DOCX</a>\n',
		encoding="utf-8",
	)
	manifest = build_lib.syllabus_model.SyllabusManifest(
		path=course_path / "syllabus.yml",
		docs_root=docs_root,
		title="Course title",
		short_name="Course",
		course_code="BIOL 000",
		term="Fall 2099",
		author="Instructor",
		language="en-US",
		course_color="#007849",
		sections=(overview_path,),
		shared_sections=(),
		lab_status="no_lab",
	)
	build_lib.syllabus_content.verify_download_links(manifest, downloads_dir)
	term_courses_path.write_text(
		'<a href="../../../downloads/Voss-BIOL_000-Fall_2099-Syllabus.pdf">PDF</a>\n',
		encoding="utf-8",
	)
	with pytest.raises(RuntimeError, match=r"fall_20xx/index\.md: missing complete-download link"):
		build_lib.syllabus_content.verify_download_links(manifest, downloads_dir)


#============================================
def test_normalize_admonitions_for_docx() -> None:
	"""Pandoc receives a portable equivalent of Material admonitions."""
	markdown = '!!! warning "Review required"\n\n    Confirm the grading plan.\n'
	portable = build_lib.syllabus_content.normalize_admonitions(markdown)
	assert portable == "**Review required**\n\nConfirm the grading plan."


#============================================
def test_docx_schedule_exam_spans_topic_and_quiz_columns(tmp_path: pathlib.Path) -> None:
	"""Word keeps the same prominent two-column exam milestone as HTML and PDF."""
	document = docx.Document()
	table = document.add_table(rows=2, cols=5)
	headers = ("Week", "Date", "Topic", "Quiz", "Due this date")
	for cell, value in zip(table.rows[0].cells, headers, strict=True):
		cell.text = value
	exam_values = ("8", "Oct 20", "", "", "Assignment 13")
	for cell, value in zip(table.rows[1].cells, exam_values, strict=True):
		cell.text = value
	exam_run = table.rows[1].cells[2].paragraphs[0].add_run("MID-TERM EXAM")
	exam_run.bold = True
	build_lib.syllabus_rendering.merge_docx_schedule_exam_cells(table, headers)
	docx_path = tmp_path / "schedule_exam.docx"
	document.save(docx_path)
	rendered_table = docx.Document(docx_path).tables[0]
	rendered_cell = rendered_table.rows[1].cells[2]
	grid_span = rendered_cell._tc.tcPr.find(docx.oxml.ns.qn("w:gridSpan"))
	observed = (
		grid_span.get(docx.oxml.ns.qn("w:val")),
		rendered_cell.text,
		rendered_table.rows[1].cells[4].text,
	)
	assert observed == ("2", "MID-TERM EXAM", "Assignment 13")


def test_remove_heading_sections_omits_web_only_policy_routes() -> None:
	"""Complete documents keep the policy branch heading without its route lists."""
	markdown = (
		"# Dr. Voss course policies\n\nShared policy introduction.\n\n"
		"## Policy topics\n\n- [Assessment](ASSESSMENT.md)\n\n"
		"## Student support\n\n- [Resources](../STUDENT_RESOURCES.md)\n"
	)
	prepared = build_lib.syllabus_content.remove_heading_sections(
		markdown,
		("Policy topics", "Student support"),
	)
	assert prepared == "# Dr. Voss course policies\n\nShared policy introduction.\n\n"


#============================================
def test_rewrite_document_links_targets_included_section(tmp_path: pathlib.Path) -> None:
	"""Source-page links become internal links when their target is included."""
	source_path = tmp_path / "course" / "ASSIGNMENTS_AND_GRADING.md"
	target_path = tmp_path / "shared" / "policies" / "ASSESSMENT.md"
	markdown = (
		"See [grade thresholds](../shared/policies/ASSESSMENT.md#grades), "
		"[grading policies](../shared/policies/ASSESSMENT.md), and "
		"[other policies](../shared/policies/OTHER.md)."
	)
	rewritten = build_lib.syllabus_content.rewrite_document_links(
		markdown,
		source_path,
		{target_path.resolve(): "assessment"},
	)
	assert rewritten == (
		"See [grade thresholds](#grades), [grading policies](#assessment), and "
		"[other policies](../shared/policies/OTHER.md)."
	)


#============================================
def test_compose_markdown_links_to_explicit_instructor_section(tmp_path: pathlib.Path) -> None:
	"""Policy routes target the instructor page listed explicitly in the manifest."""
	term_path = tmp_path / "fall_20xx"
	course_path = term_path / "course"
	shared_path = term_path / "shared"
	policy_path = shared_path / "policies"
	course_path.mkdir(parents=True)
	policy_path.mkdir(parents=True)
	index_path = course_path / "index.md"
	details_path = course_path / "COURSE_DETAILS.md"
	policies_path = policy_path / "index.md"
	instructor_route_path = shared_path / "INSTRUCTOR_INFORMATION.md"
	write_section(index_path, "Course title")
	details_path.write_text(
		"# Course information\n\n"
		"<!-- lab attendance from syllabus.yml -->\n",
		encoding="utf-8",
	)
	policies_path.write_text(
		"# Policies\n\nSee [Instructor information](../INSTRUCTOR_INFORMATION.md).\n",
		encoding="utf-8",
	)
	instructor_route_path.write_text(
		"# Instructor information\n\nContact details.\n",
		encoding="utf-8",
	)
	manifest = build_lib.syllabus_model.SyllabusManifest(
		path=course_path / "syllabus.yml",
		docs_root=tmp_path,
		title="Course title",
		short_name="Course",
		course_code="BIOL 000",
		term="Fall 2099",
		author="Instructor",
		language="en-US",
		course_color="#007849",
		sections=(index_path, details_path, instructor_route_path),
		shared_sections=(policies_path,),
		lab_status="no_lab",
	)
	combined = build_lib.syllabus_content.compose_markdown(manifest)
	assert "[Instructor information](#instructor-information)" in combined
	assert "(../INSTRUCTOR_INFORMATION.md)" not in combined
	assert combined.count("Contact details.") == 1


#============================================
def test_instructor_page_titles_derive_from_source_names(tmp_path: pathlib.Path) -> None:
	"""Document navigation uses Title Case page names rather than student headings."""
	assert build_lib.syllabus_content.get_instructor_page_title(
		tmp_path / "ASSIGNMENTS_AND_GRADING.md"
	) == "Assignments and Grading"
	assert build_lib.syllabus_content.get_instructor_page_title(
		tmp_path / "policies" / "index.md"
	) == "Course Policies"


#============================================
def test_markdown_html_uses_site_extension_stack(tmp_path: pathlib.Path) -> None:
	"""PDF HTML preserves native admonitions and explicit heading anchors."""
	manifest = build_lib.syllabus_model.SyllabusManifest(
		path=tmp_path / "syllabus.yml",
		docs_root=tmp_path,
		title="Course title",
		short_name="Course",
		course_code="BIOL 000",
		term="Fall 2099",
		author="Instructor",
		language="en-US",
		course_color="#007849",
		sections=(),
		shared_sections=(),
		lab_status="no_lab",
	)
	stylesheet_path = tmp_path / "print.css"
	stylesheet_path.write_text("body { color: black; }\n", encoding="utf-8")
	html_path = tmp_path / "syllabus.html"
	build_lib.syllabus_rendering.run_markdown_html(
		'## Course overview {#course-overview}\n\n!!! warning "Review"\n\n    Check this.\n\n'
		"| Field | Value |\n| --- | ---: |\n| Points | 10 |\n",
		html_path,
		stylesheet_path,
		manifest,
		("admonition", "attr_list", "tables"),
		{},
	)
	html_text = html_path.read_text(encoding="utf-8")
	assert '<html lang="en-US">' in html_text
	assert f'<base href="{tmp_path.as_uri()}/">' in html_text
	assert 'style="--syllabus-page-accent: #007849"' in html_text
	assert '<h1 data-course-short-name="Course">BIOL 000: Course title</h1>' in html_text
	assert (
		'<h2 id="course-overview">Course overview</h2>\n'
		'<div class="admonition warning">\n<p class="admonition-title">Review</p>'
	) in html_text
	assert html_text.count("<table>") == 1


#============================================
def test_course_theme_rejects_unsafe_css_values(tmp_path: pathlib.Path) -> None:
	"""Course metadata accepts only auditable six-digit hex colors."""
	metadata_path = tmp_path / ".meta.yml"
	metadata_path.write_text(
		'course_color: "red; display: none"\ncourse_color_dark: "#73c167"\n',
		encoding="utf-8",
	)
	with pytest.raises(ValueError, match="course_color must be a six-digit hex color"):
		build_lib.syllabus_model.load_course_theme(metadata_path)


#============================================
def test_secret_scan_rejects_meeting_credentials() -> None:
	"""Credential-shaped public text fails closed."""
	credential_fixture = "".join(
		(
			"Join at https://example.",
			"zoom.",
			"us/j/123456789?",
			"pwd",
			"=synthetic-value",
		)
	)
	with pytest.raises(ValueError, match="prohibited public credential pattern"):
		build_lib.syllabus_content.scan_text_for_secrets(
			credential_fixture,
			"inline test",
		)


#============================================
def test_source_scan_rejects_line_breaking_control_characters() -> None:
	"""Invisible controls cannot silently turn a Markdown table into prose."""
	with pytest.raises(ValueError, match="prohibited control character: U\\+000B"):
		build_lib.syllabus_content.scan_text_for_prohibited_controls(
			"| Included in\vtotal points |",
			"inline test",
		)


#============================================
def test_markdown_table_validation_requires_consistent_columns() -> None:
	"""Every Markdown table has named headers and a rectangular row structure."""
	valid_markdown = "| Absence type | Score |\n| --- | ---: |\n| First communicated | N/A |\n"
	build_lib.syllabus_content.validate_markdown_tables(valid_markdown, "inline test")
	assert build_lib.syllabus_content.count_markdown_tables(valid_markdown) == 1
	markdown = "| Absence type | Score |\n| --- | ---: |\n| First communicated |\n"
	with pytest.raises(ValueError, match="inconsistent table columns"):
		build_lib.syllabus_content.validate_markdown_tables(markdown, "inline test")


#============================================
def write_section(path: pathlib.Path, heading: str) -> None:
	"""Write one minimal inline section for composition testing."""
	path.write_text(f"# {heading}\n\n{heading} body.\n", encoding="utf-8")
	return None


#============================================
def test_learning_framework_requires_all_four_ordered_sections(tmp_path: pathlib.Path) -> None:
	"""Every course preserves the distinct syllabus learning statements in order."""
	framework_path = tmp_path / "COURSE_LEARNING_FRAMEWORK.md"
	framework_path.write_text(
		"# Learning Objectives, Outcomes, and Goals\n\n"
		"## Roosevelt learning goals\n\n- Communication.\n\n"
		"## Learning Objectives:\n\n"
		"Students completing this course will have achieved:\n\n- Experience.\n\n"
		"## Course Learning Outcomes:\n\n"
		"Students completing this course will be able to:\n\n- Apply knowledge.\n\n"
		"## Learning Goals:\n\n"
		"Overall, this course aims to accomplish:\n\n- Growth.\n",
		encoding="utf-8",
	)
	build_lib.syllabus_content.validate_course_learning_framework((framework_path,), tmp_path)


#============================================
def test_compose_markdown_appends_policy_and_resources_once(tmp_path: pathlib.Path) -> None:
	"""The manifest order keeps policies and resources separate and non-duplicated."""
	index_path = tmp_path / "index.md"
	policy_path = tmp_path / "shared" / "policies" / "index.md"
	resource_path = tmp_path / "shared" / "STUDENT_RESOURCES.md"
	policy_path.parent.mkdir(parents=True)
	write_section(index_path, "Course title")
	write_section(policy_path, "Policies")
	write_section(resource_path, "Student resources")
	manifest = build_lib.syllabus_model.SyllabusManifest(
		path=tmp_path / "syllabus.yml",
		docs_root=tmp_path,
		title="Course title",
		short_name="Course",
		course_code="BIOL 000",
		term="Fall 2099",
		author="Instructor",
		language="en-US",
		course_color="#007849",
		sections=(index_path,),
		shared_sections=(policy_path, resource_path),
		lab_status="no_lab",
	)
	combined = build_lib.syllabus_content.compose_markdown(manifest)
	assert combined.count("## Policies {#policies}") == 1
	assert combined.index("## Policies") < combined.index("## Student resources")


#============================================
def test_manifest_assessments_control_coursework_content(
	tmp_path: pathlib.Path,
) -> None:
	"""Manifest selection determines assessment content and order."""
	docs_root = tmp_path / "site_docs"
	term_path = docs_root / "fall_20xx"
	course_path = term_path / "course"
	fragment_path = term_path / "shared" / "fragments" / "assessments"
	course_path.mkdir(parents=True)
	fragment_path.mkdir(parents=True)
	index_path = course_path / "index.md"
	coursework_path = course_path / "ASSIGNMENTS_AND_GRADING.md"
	assignments_path = fragment_path / "ASSIGNMENTS.md"
	exams_path = fragment_path / "EXAMS.md"
	write_section(index_path, "Course title")
	coursework_path.write_text(
		"# Coursework and grades\n\n<!-- assessments from syllabus.yml -->\n",
		encoding="utf-8",
	)
	assignments_path.write_text("## Assignments\n\nAssignment rules.\n", encoding="utf-8")
	exams_path.write_text("## Exams\n\nExam rules.\n", encoding="utf-8")
	manifest = build_lib.syllabus_model.SyllabusManifest(
		path=course_path / "syllabus.yml",
		docs_root=docs_root,
		title="Course title",
		short_name="Course",
		course_code="BIOL 000",
		term="Fall 2099",
		author="Instructor",
		language="en-US",
		course_color="#007849",
		sections=(index_path, coursework_path),
		shared_sections=(),
		lab_status="no_lab",
		assessment_sections=(
			build_lib.syllabus_model.AssessmentSection(root_fragment=exams_path),
			build_lib.syllabus_model.AssessmentSection(root_fragment=assignments_path),
		),
	)
	combined = build_lib.syllabus_content.compose_markdown(manifest)
	assert combined.index("\n### Exams\n") < combined.index("\n### Assignments\n")
	assert "assessments from syllabus.yml" not in combined


#============================================
def test_assessment_fragment_requires_level_two_root(tmp_path: pathlib.Path) -> None:
	"""Assessment fragments fail before a section root can be visually flattened."""
	fragment_path = tmp_path / "ASSIGNMENTS.md"
	fragment_path.write_text("### Assignment details\n\nRules.\n", encoding="utf-8")
	with pytest.raises(ValueError, match="must begin with a level-two heading"):
		build_lib.syllabus_content.validate_assessment_section_fragment(fragment_path)


#============================================
def test_assessment_fragment_rejects_heading_level_skips(tmp_path: pathlib.Path) -> None:
	"""Assessment subsections remain H3 children of their H2 section."""
	fragment_path = tmp_path / "ASSIGNMENTS.md"
	fragment_path.write_text(
		"## Assignment details\n\n#### Practice until it makes sense\n",
		encoding="utf-8",
	)
	with pytest.raises(ValueError, match="allow only H2 and H3"):
		build_lib.syllabus_content.validate_assessment_section_fragment(fragment_path)


#============================================
def test_assessment_topic_requires_level_three_root(tmp_path: pathlib.Path) -> None:
	"""Selected policy topics remain H3 children of their shared H2 section."""
	fragment_path = tmp_path / "TECHNOLOGY_INTERRUPTION_ASSIGNMENTS.md"
	fragment_path.write_text("## Assignments\n\nRecovery.\n", encoding="utf-8")
	with pytest.raises(ValueError, match="must begin with a level-three heading"):
		build_lib.syllabus_content.validate_assessment_topic_fragment(fragment_path)


#============================================
def test_assessment_manifest_rejects_unknown_category(tmp_path: pathlib.Path) -> None:
	"""The manifest accepts only Dr. Voss's four assessment categories."""
	manifest_path = tmp_path / "fall_20xx" / "course" / "syllabus.yml"
	manifest_path.parent.mkdir(parents=True)
	with pytest.raises(ValueError, match="unsupported assessments"):
		build_lib.syllabus_model.resolve_assessment_sections(
			{"assessments": ["lab_practicals"]},
			manifest_path,
			tmp_path,
		)


#============================================
def test_coursework_requires_one_assessment_marker(tmp_path: pathlib.Path) -> None:
	"""Configured coursework fails rather than silently omitting its assessments."""
	coursework_path = tmp_path / "ASSIGNMENTS_AND_GRADING.md"
	manifest = build_lib.syllabus_model.SyllabusManifest(
		path=tmp_path / "syllabus.yml",
		docs_root=tmp_path,
		title="Course title",
		short_name="Course",
		course_code="BIOL 000",
		term="Fall 2099",
		author="Instructor",
		language="en-US",
		course_color="#007849",
		sections=(coursework_path,),
		shared_sections=(),
		lab_status="no_lab",
		assessment_sections=(
			build_lib.syllabus_model.AssessmentSection(
				root_fragment=tmp_path / "ASSIGNMENTS.md"
			),
		),
	)
	with pytest.raises(ValueError, match="expected exactly one assessment fragment marker"):
		build_lib.syllabus_content.apply_assessment_fragments(
			"# Coursework and grades\n",
			coursework_path,
			manifest,
		)


#============================================
def test_discussion_manifest_rejects_unknown_mode(tmp_path: pathlib.Path) -> None:
	"""The manifest accepts only Dr. Voss's three discussion modes."""
	manifest_path = tmp_path / "fall_20xx" / "course" / "syllabus.yml"
	manifest_path.parent.mkdir(parents=True)
	with pytest.raises(ValueError, match="unsupported discussion mode"):
		build_lib.syllabus_model.resolve_discussion_fragments(
			{"discussion": "generic_participation"},
			manifest_path,
			tmp_path,
		)


#============================================
def test_discussion_marker_materializes_selected_fragments(tmp_path: pathlib.Path) -> None:
	"""One course discussion page receives only its selected mode fragments."""
	discussion_path = tmp_path / "DISCUSSION_MARKS.md"
	mode_path = tmp_path / "FACE_TO_FACE.md"
	common_path = tmp_path / "COMMON.md"
	manifest = build_lib.syllabus_model.SyllabusManifest(
		path=tmp_path / "syllabus.yml",
		docs_root=tmp_path,
		title="Course title",
		short_name="Course",
		course_code="BIOL 000",
		term="Fall 2099",
		author="Instructor",
		language="en-US",
		course_color="#007849",
		sections=(discussion_path,),
		shared_sections=(),
		lab_status="no_lab",
		discussion_fragments=(mode_path, common_path),
	)
	selected = build_lib.syllabus_content.apply_discussion_fragments(
		"# Discussion marks\n\n<!-- discussion from syllabus.yml -->\n",
		discussion_path,
		manifest,
	)
	assert selected.index("FACE_TO_FACE.md") < selected.index("COMMON.md")
	assert "discussion from syllabus.yml" not in selected


#============================================
def test_lab_manifest_accepts_closed_status_and_rejects_unknown(
	tmp_path: pathlib.Path,
) -> None:
	"""Lab policy inclusion accepts only the documented binary course state."""
	manifest_path = tmp_path / "fall_20xx" / "course" / "syllabus.yml"
	manifest_path.parent.mkdir(parents=True)
	fragment_path = (
		tmp_path
		/ "fall_20xx"
		/ "shared"
		/ "fragments"
		/ "labs"
		/ "LAB_ATTENDANCE.md"
	)
	fragment_path.parent.mkdir(parents=True)
	fragment_path.write_text("## Lab attendance\n", encoding="utf-8")
	lab_status, fragments = build_lib.syllabus_model.resolve_lab_fragments(
		{"lab_status": "has_lab"},
		manifest_path,
		tmp_path,
	)
	assert (lab_status, fragments) == ("has_lab", (fragment_path,))
	with pytest.raises(ValueError, match="lab_status must be one of"):
		build_lib.syllabus_model.resolve_lab_fragments(
			{"lab_status": "maybe_lab"},
			manifest_path,
			tmp_path,
		)


#============================================
def test_lab_status_controls_course_details_content(tmp_path: pathlib.Path) -> None:
	"""Only a syllabus declaring a lab receives the canonical attendance fragment."""
	docs_root = tmp_path / "site_docs"
	term_path = docs_root / "fall_20xx"
	course_path = term_path / "course"
	fragment_path = term_path / "shared" / "fragments" / "labs" / "LAB_ATTENDANCE.md"
	course_path.mkdir(parents=True)
	fragment_path.parent.mkdir(parents=True)
	index_path = course_path / "index.md"
	details_path = course_path / "COURSE_DETAILS.md"
	write_section(index_path, "Course title")
	details_path.write_text(
		"# Course information\n\n"
		"General course details.\n\n"
		"<!-- lab attendance from syllabus.yml -->\n",
		encoding="utf-8",
	)
	fragment_path.write_text(
		"## Lab attendance\n\nLab-only attendance policy.\n",
		encoding="utf-8",
	)
	common_fields = {
		"path": course_path / "syllabus.yml",
		"docs_root": docs_root,
		"title": "Course title",
		"short_name": "Course",
		"course_code": "BIOL 000",
		"term": "Fall 2099",
		"author": "Instructor",
		"language": "en-US",
		"course_color": "#007849",
		"sections": (index_path, details_path),
		"shared_sections": (),
	}
	no_lab_manifest = build_lib.syllabus_model.SyllabusManifest(
		**common_fields,
		lab_status="no_lab",
	)
	has_lab_manifest = build_lib.syllabus_model.SyllabusManifest(
		**common_fields,
		lab_status="has_lab",
		lab_fragments=(fragment_path,),
	)
	no_lab_markdown = build_lib.syllabus_content.compose_markdown(no_lab_manifest)
	has_lab_markdown = build_lib.syllabus_content.compose_markdown(has_lab_manifest)
	assert "General course details." in no_lab_markdown
	assert "Lab-only attendance policy." not in no_lab_markdown
	assert "Lab-only attendance policy." in has_lab_markdown
	assert "lab attendance from syllabus.yml" not in has_lab_markdown


#============================================
def test_publish_downloads_replaces_the_managed_set(tmp_path: pathlib.Path) -> None:
	"""Validated staged downloads replace current files and remove obsolete siblings."""
	staged_dir = tmp_path / "staged"
	downloads_dir = tmp_path / "downloads"
	staged_dir.mkdir()
	downloads_dir.mkdir()
	(staged_dir / "COURSE.pdf").write_text("new PDF", encoding="utf-8")
	(staged_dir / "COURSE.docx").write_text("new DOCX", encoding="utf-8")
	(downloads_dir / "COURSE.pdf").write_text("old PDF", encoding="utf-8")
	(downloads_dir / "STALE.docx").write_text("obsolete", encoding="utf-8")
	(downloads_dir / "README.txt").write_text("preserve", encoding="utf-8")
	build_lib.syllabus_rendering.publish_downloads(
		staged_dir,
		downloads_dir,
		{"COURSE.docx", "COURSE.pdf"},
	)
	managed_state = (
		(downloads_dir / "COURSE.pdf").read_text(encoding="utf-8"),
		(downloads_dir / "COURSE.docx").read_text(encoding="utf-8"),
		(downloads_dir / "STALE.docx").exists(),
	)
	assert managed_state == ("new PDF", "new DOCX", False)
	assert (downloads_dir / "README.txt").read_text(encoding="utf-8") == "preserve"


#============================================
def test_publish_downloads_rejects_an_incomplete_stage(tmp_path: pathlib.Path) -> None:
	"""A partial staged build cannot replace an existing published download."""
	staged_dir = tmp_path / "staged"
	downloads_dir = tmp_path / "downloads"
	staged_dir.mkdir()
	downloads_dir.mkdir()
	(staged_dir / "COURSE.pdf").write_text("new PDF", encoding="utf-8")
	current_path = downloads_dir / "COURSE.pdf"
	current_path.write_text("current PDF", encoding="utf-8")
	with pytest.raises(RuntimeError, match="Staged downloads do not match"):
		build_lib.syllabus_rendering.publish_downloads(
			staged_dir,
			downloads_dir,
			{"COURSE.docx", "COURSE.pdf"},
		)
	assert current_path.read_text(encoding="utf-8") == "current PDF"
