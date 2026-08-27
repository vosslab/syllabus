"""Build every syllabus output and verify Markdown include parity."""

# Standard Library
import pathlib
import zipfile
import subprocess
import html.parser

# PIP3 modules
import docx


INCLUDE_MARKER = "--8<--"


class HTMLTextParser(html.parser.HTMLParser):
	"""Collect text nodes from built HTML."""

	def __init__(self) -> None:
		super().__init__()
		self.parts: list[str] = []

	def handle_data(self, data: str) -> None:
		"""Collect one HTML text node."""
		self.parts.append(data)
		return None

	def get_text(self) -> str:
		"""Return normalized text from every collected node."""
		html_text = normalize_text(" ".join(self.parts))
		return html_text


#============================================
def get_repo_root() -> pathlib.Path:
	"""Return the repository root reported by Git."""
	completed = subprocess.run(
		["git", "rev-parse", "--show-toplevel"],
		check=True,
		capture_output=True,
		text=True,
	)
	repo_root = pathlib.Path(completed.stdout.strip())
	return repo_root


#============================================
def normalize_text(text_value: str) -> str:
	"""Collapse renderer-specific whitespace for cross-format comparisons."""
	normalized = " ".join(text_value.split())
	return normalized


#============================================
def read_site_text(site_root: pathlib.Path) -> str:
	"""Extract normalized text nodes from every built HTML page."""
	parts = []
	for html_path in sorted(site_root.rglob("*.html")):
		parser = HTMLTextParser()
		parser.feed(html_path.read_text(encoding="utf-8"))
		parts.append(parser.get_text())
	combined = normalize_text(" ".join(parts))
	return combined


#============================================
def read_docx_text(docx_path: pathlib.Path) -> str:
	"""Extract normalized paragraph and table text from one DOCX."""
	document = docx.Document(docx_path)
	parts = [paragraph.text for paragraph in document.paragraphs]
	for table in document.tables:
		for row in table.rows:
			for cell in row.cells:
				parts.extend(paragraph.text for paragraph in cell.paragraphs)
	combined = normalize_text(" ".join(parts))
	return combined


#============================================
def read_pdf_text(pdf_path: pathlib.Path) -> str:
	"""Extract normalized selectable text from one PDF through Poppler."""
	completed = subprocess.run(
		["pdftotext", str(pdf_path), "-"],
		check=True,
		capture_output=True,
		text=True,
	)
	text_content = normalize_text(completed.stdout)
	return text_content


#============================================
def find_marker_files(root_path: pathlib.Path) -> list[str]:
	"""Return files whose built bytes still contain an include marker."""
	marker = INCLUDE_MARKER.encode("ascii")
	matches = []
	for candidate in sorted(root_path.rglob("*")):
		if candidate.is_file() and marker in candidate.read_bytes():
			matches.append(candidate.relative_to(root_path).as_posix())
	return matches


#============================================
def find_docx_marker_files(downloads_root: pathlib.Path) -> list[str]:
	"""Return DOCX packages whose XML text still contains an include marker."""
	marker = INCLUDE_MARKER.encode("ascii")
	matches = []
	for docx_path in sorted(downloads_root.glob("*.docx")):
		with zipfile.ZipFile(docx_path) as archive:
			contains_marker = any(
				member_name.endswith(".xml") and marker in archive.read(member_name)
				for member_name in archive.namelist()
			)
		if contains_marker:
			matches.append(docx_path.name)
	return matches


#============================================
def select_contact_sentence(fragment_path: pathlib.Path) -> str:
	"""Select one source sentence that should reach every output corpus."""
	fragment_text = fragment_path.read_text(encoding="utf-8")
	paragraph = fragment_text.split("## How to contact me", 1)[1].strip().split("\n\n", 1)[0]
	sentence = normalize_text(paragraph).split(". ", 1)[0] + "."
	return sentence


#============================================
def select_generated_event(fragment_path: pathlib.Path) -> str:
	"""Select one current event value that should reach every artifact corpus."""
	for line in fragment_path.read_text(encoding="utf-8").splitlines():
		if not line.startswith("|") or line.startswith(("| Date |", "| --- |")):
			continue
		cells = [cell.strip() for cell in line.strip("|").split("|")]
		event_name = cells[1]
		return event_name
	raise ValueError(f"{fragment_path}: generated dates contain no event rows")


#============================================
def main() -> None:
	"""Run the production export E2E and verify cross-format include evidence."""
	repo_root = get_repo_root()
	print("\n=== Build complete export and strict site ===")
	subprocess.run(
		["bash", "tests/e2e/e2e_syllabus_export.sh"],
		cwd=repo_root,
		check=True,
	)

	print("\n=== Inspect include expansion across artifacts ===")
	site_root = repo_root / "site"
	downloads_root = repo_root / "site_docs" / "downloads"
	site_marker_files = find_marker_files(site_root)
	docx_marker_files = find_docx_marker_files(downloads_root)
	pdf_marker_files = []
	docx_text_parts = []
	pdf_text_parts = []
	docx_generated_missing = []
	pdf_generated_missing = []
	generated_event = select_generated_event(
		repo_root / "site_docs" / "generated" / "FALL_2026_IMPORTANT_DATES.md"
	)
	for docx_path in sorted(downloads_root.glob("*.docx")):
		docx_text = read_docx_text(docx_path)
		docx_text_parts.append(docx_text)
		if generated_event not in docx_text:
			docx_generated_missing.append(docx_path.name)
	for pdf_path in sorted(downloads_root.glob("*.pdf")):
		pdf_text = read_pdf_text(pdf_path)
		pdf_text_parts.append(pdf_text)
		if INCLUDE_MARKER in pdf_text:
			pdf_marker_files.append(pdf_path.name)
		if generated_event not in pdf_text:
			pdf_generated_missing.append(pdf_path.name)
	assert not site_marker_files, f"Unexpanded site includes: {site_marker_files}"
	assert not docx_marker_files, f"Unexpanded DOCX includes: {docx_marker_files}"
	assert not pdf_marker_files, f"Unexpanded PDF includes: {pdf_marker_files}"

	contact_sentence = select_contact_sentence(
		repo_root
		/ "site_docs"
		/ "fall_2026"
		/ "shared"
		/ "fragments"
		/ "INSTRUCTOR_CONTACT_DETAILS.md"
	)
	site_text = read_site_text(site_root)
	docx_text = normalize_text(" ".join(docx_text_parts))
	pdf_text = normalize_text(" ".join(pdf_text_parts))
	assert contact_sentence in site_text, "Shared fragment is absent from the website corpus"
	assert contact_sentence in docx_text, "Shared fragment is absent from the DOCX corpus"
	assert contact_sentence in pdf_text, "Shared fragment is absent from the PDF corpus"
	assert generated_event in site_text, "Generated fragment is absent from the website corpus"
	assert not docx_generated_missing, f"Generated fragment is absent from: {docx_generated_missing}"
	assert not pdf_generated_missing, f"Generated fragment is absent from: {pdf_generated_missing}"
	print("PASS: include expansion is consistent across HTML, DOCX, and PDF artifacts")
	return None


if __name__ == "__main__":
	main()
